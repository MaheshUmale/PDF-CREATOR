import os
import re
import gdown
from PyPDF2 import PdfMerger
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT


# def extract_hyperlinks_from_word(file_path):
#     document = Document(file_path)
#     hyperlinks = []

#     for paragraph in document.paragraphs:
#         if not paragraph.hyperlinks:
#             continue
#         for hyperlink in paragraph.hyperlinks:
#             print(hyperlink.text)
#             print(hyperlink.address)
#             hyperlinks.append(hyperlink.text+":"+hyperlink.address)
#     # Also check document-level relationships for hyperlinks (e.g., in headers/footers)
#     for rel in document.part.rels:
#         if document.part.rels[rel].reltype == RT.HYPERLINK:
#             hyperlinks.append(document.part.rels[rel].target_ref)
#     return hyperlinks

# Example usage:
word_document_path = 'D:\\py_code_workspace\\PDF CREATOR\\index.docx'
# extracted_links = extract_hyperlinks_from_word(word_document_path)
# for link in extracted_links:
#     print(link)
    

# --- SCRIPT SETUP ---
# NOTE: This script requires the following libraries.
# You can install them by running:
# pip install gdown PyPDF2 python-docx

# --- CONFIGURATION ---
# Set the path to your Word document that contains the links to all your chapters.
# This file should be placed in the same directory as this script.
WORD_DOC_PATH = 'D:\\py_code_workspace\\PDF CREATOR\\index.docx'

# Regular expression to extract the document ID from a Google Docs URL
DOC_ID_PATTERN = r'https://docs.google.com/document/d/([a-zA-Z0-9_-]+)'

# --- HELPER FUNCTIONS ---

def get_doc_links_from_word(file_path):
    """
    Parses a Word document to find all hyperlinks that are Google Docs links.
    Returns a list of dictionaries, each containing a 'title' and 'id'.
    """
    if not os.path.exists(file_path):
        print(f"Error: The file '{file_path}' was not found.")
        return []

    print(f"\nParsing '{file_path}' to find Google Docs links...")
    hyperlinks = []
    
    try:
        document = Document(file_path)
    except Exception as e:
        print(f"Error opening document: {e}")
        return []

    # Iterate through all paragraphs and their hyperlinks
    
    for paragraph in document.paragraphs:
        if not paragraph.hyperlinks:
            continue
        for hyperlink in paragraph.hyperlinks:
            print(hyperlink.text)
            print(hyperlink.address)
            hyperlinks.append(hyperlink.text+":"+hyperlink.address)
    
            link_text = hyperlink.text.strip()
            print(link_text)
            url = hyperlink.address
            match = re.search(DOC_ID_PATTERN, url)
            if match:
                doc_id = match.group(1)
                # Use the link text as the title, or a default if it's empty
                title = link_text if link_text else f"Untitled Document ({doc_id[:8]})"
                hyperlinks.append({'id': doc_id, 'title': title})

    # Also check document-level relationships (e.g., in headers/footers)
    for rel in document.part.rels:
        if document.part.rels[rel].reltype == RT.HYPERLINK:
            url = document.part.rels[rel].target_ref
            match = re.search(DOC_ID_PATTERN, url)
            if match:
                doc_id = match.group(1)
                title = f"Untitled Document ({doc_id[:8]})"
                hyperlinks.append({'id': doc_id, 'title': title})

    # Deduplicate links in case the same one appears multiple times
    seen = set()
    unique_links = []
    for link in hyperlinks:
        if link['id'] not in seen:
            seen.add(link['id'])
            unique_links.append(link)

    print(f"Found {len(unique_links)} Google Docs links.")
    return unique_links

def download_doc_as_pdf(doc_id, doc_title, output_dir='pdfs'):
    """
    Downloads a single Google Doc as a PDF file using gdown.
    Returns the path to the downloaded PDF.
    """
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # Use the unique document ID to ensure a unique filename
    file_name = f"{doc_id}.pdf"
    file_path = os.path.join(output_dir, file_name)
    
    print(f"\nDownloading '{doc_title}' ({doc_id}) as PDF using gdown...")
    
    # gdown automatically handles the conversion of a Google Doc to PDF.
    gdown.download(id=doc_id, output=file_path, quiet=False, fuzzy=True)
    
    print(f"  Successfully downloaded to '{file_path}'")
    return file_path

def merge_pdfs(pdf_paths, output_filename='Ebook_Combined.pdf'):
    """Merges a list of PDF files into a single output PDF."""
    print(f"\nMerging {len(pdf_paths)} PDFs into '{output_filename}'...")
    merger = PdfMerger()
    
    for pdf_path in pdf_paths:
        try:
            merger.append(pdf_path)
            print(f"  Appended {os.path.basename(pdf_path)}")
        except Exception as e:
            print(f"  Error appending {pdf_path}: {e}")
            
    merger.write(output_filename)
    merger.close()
    
    print(f"\nSuccessfully created the final ebook: '{output_filename}'")
    
# --- MAIN SCRIPT EXECUTION ---

if __name__ == '__main__':
    # Step 1: Get the list of chapters from the Word document
    chapters = get_doc_links_from_word(WORD_DOC_PATH)

    if not chapters:
        print("No chapters found. Please ensure your 'index.docx' exists and contains valid Google Docs links.")
    else:
        # Step 2: Download each chapter as a PDF
        downloaded_pdfs = []
        for chapter in chapters:
            try:
                pdf_path = download_doc_as_pdf(chapter['id'], chapter['title'])
                downloaded_pdfs.append(pdf_path)
            except Exception as e:
                print(f"Error downloading {chapter['title']}: {e}")
        
        if downloaded_pdfs:
            # Step 3: Merge the downloaded PDFs into a single file
            merge_pdfs(downloaded_pdfs)
